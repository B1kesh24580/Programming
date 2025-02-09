import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Global variable for the Word document
doc = None

def open_file():
    """Open a Word file and display its content."""
    global doc
    file_path = filedialog.askopenfilename(
        title="Select a Word File",
        filetypes=[("Word Documents", "*.docx")]
    )
    if file_path:
        try:
            doc = Document(file_path)
            display_content()
            messagebox.showinfo("Success", "File loaded successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open file: {e}")
    else:
        messagebox.showwarning("No File", "No file selected.")

def display_content():
    """Display the content of the document in the text area."""
    if doc:
        text_area.delete("1.0", tk.END)
        for i, paragraph in enumerate(doc.paragraphs):
            text_area.insert(tk.END, f"{i + 1}: {paragraph.text}\n")
    else:
        messagebox.showwarning("No Document", "No document is loaded.")

def modify_paragraph():
    """Modify the specified paragraph."""
    if doc:
        try:
            para_num = int(paragraph_number_entry.get()) - 1
            new_text = new_text_entry.get()
            alignment_choice = alignment_var.get()

            if 0 <= para_num < len(doc.paragraphs):
                paragraph = doc.paragraphs[para_num]
                paragraph.text = new_text

                # Apply alignment
                if alignment_choice == "Center":
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif alignment_choice == "Right":
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                display_content()
                messagebox.showinfo("Success", "Paragraph modified successfully!")
            else:
                messagebox.showerror("Error", "Invalid paragraph number.")
        except ValueError:
            messagebox.showerror("Error", "Please enter a valid paragraph number.")
    else:
        messagebox.showwarning("No Document", "No document is loaded.")

def add_paragraph():
    """Add a new paragraph at the end of the document."""
    if doc:
        new_text = new_text_entry.get()
        doc.add_paragraph(new_text)
        display_content()
        messagebox.showinfo("Success", "New paragraph added successfully!")
    else:
        messagebox.showwarning("No Document", "No document is loaded.")

def save_file():
    """Save the modified document."""
    if doc:
        save_path = filedialog.asksaveasfilename(
            title="Save the Word File",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if save_path:
            try:
                doc.save(save_path)
                messagebox.showinfo("Success", f"Document saved successfully as '{save_path}'.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save file: {e}")
        else:
            messagebox.showwarning("No Save Location", "No save location selected.")
    else:
        messagebox.showwarning("No Document", "No document is loaded.")

# Create the main Tkinter window
root = tk.Tk()
root.title("Word File Editor")

# GUI Widgets
open_button = tk.Button(root, text="Open File", command=open_file)
open_button.pack(pady=5)

text_area = tk.Text(root, width=80, height=20)
text_area.pack(pady=5)

paragraph_number_label = tk.Label(root, text="Paragraph Number:")
paragraph_number_label.pack(pady=2)
paragraph_number_entry = tk.Entry(root, width=10)
paragraph_number_entry.pack(pady=2)

new_text_label = tk.Label(root, text="New Text:")
new_text_label.pack(pady=2)
new_text_entry = tk.Entry(root, width=50)
new_text_entry.pack(pady=2)

alignment_label = tk.Label(root, text="Alignment:")
alignment_label.pack(pady=2)
alignment_var = tk.StringVar(value="Left")
alignment_menu = tk.OptionMenu(root, alignment_var, "Left", "Center", "Right")
alignment_menu.pack(pady=2)

modify_button = tk.Button(root, text="Modify Paragraph", command=modify_paragraph)
modify_button.pack(pady=5)

add_button = tk.Button(root, text="Add Paragraph", command=add_paragraph)
add_button.pack(pady=5)

save_button = tk.Button(root, text="Save File", command=save_file)
save_button.pack(pady=5)

# Start the GUI event loop
root.mainloop()
