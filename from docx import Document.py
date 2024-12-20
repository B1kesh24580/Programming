from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading('Document Title', level=1)

# Add a paragraph
paragraph = doc.add_paragraph('This is a sample paragraph. You can change the font, increase font size, bold, italicize, and change indentations.')

# Change font and size
run = paragraph.add_run(' This text has a different font and size.')
run.font.name = 'Arial'
run.font.size = Pt(14)

# Bold and Italicize text
bold_run = paragraph.add_run(' This text is bold.')
bold_run.bold = True

italic_run = paragraph.add_run(' This text is italic.')
italic_run.italic = True

# Change line indent and hanging indent
paragraph.paragraph_format.left_indent = Pt(36)  # Left indent
paragraph.paragraph_format.first_line_indent = Pt(-18)  # Hanging indent

# Add a new paragraph with different line spacing
new_paragraph = doc.add_paragraph('This paragraph has different line spacing.')
new_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# Add page numbers
section = doc.sections[0]
footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = "Page "
footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer_paragraph.add_run("G:\BIkesh M 2081\Solution\SEE Self study Book\Self Study Book Economic\Exam Tips for economic.docx")
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading('Document Title', level=1)

# Add a paragraph
paragraph = doc.add_paragraph('This is a sample paragraph. You can change the font, increase font size, bold, italicize, and change indentations.')

# Change font and size
run = paragraph.add_run(' This text has a different font and size.')
run.font.name = 'Arial'
run.font.size = Pt(14)

# Bold and Italicize text
bold_run = paragraph.add_run(' This text is bold.')
bold_run.bold = True

italic_run = paragraph.add_run(' This text is italic.')
italic_run.italic = True

# Change line indent and hanging indent
paragraph.paragraph_format.left_indent = Pt(36)  # Left indent
paragraph.paragraph_format.first_line_indent = Pt(-18)  # Hanging indent

# Add a new paragraph with different line spacing
new_paragraph = doc.add_paragraph('This paragraph has different line spacing.')
new_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# Add page numbers
section = doc.sections[0]
footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = "Page "
footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# footer_paragraph.add_run("").add_field('PAGE')

