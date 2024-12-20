from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load the document
doc = Document('G:\BIkesh M 2081\Solution\SEE Self study Book\Self Study Book Economic\Exam Tips for economic.docx')

# Change the style of the first paragraph
paragraph = doc.paragraphs[0]
run = paragraph.runs[0]
run.font.name = 'Arial'
run.font.size = Pt(12)
run.bold = True

# Center align the second paragraph
paragraph = doc.paragraphs[1]
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add a new paragraph
doc.add_paragraph("Hello this is the test subject")

# Save the changes
doc.save('G:\BIkesh M 2081\Solution\SEE Self study Book\Self Study Book Economic')